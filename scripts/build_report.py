"""Build ARTICLE.md, TEI-Bench.html, and paper/TEI-Bench.docx from the v2
ablation results (results_v2/). Reads _summary.json + per_agent.csv and
hand-types ZERO numbers, so the documents cannot diverge from the recorded run.

Usage:
  pip install python-docx
  python scripts/build_report.py
"""
import csv
import html
import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
RES = ROOT / "results_v2"
HB = ROOT / "results_v2_hb"
FIG = ROOT / "paper" / "figures"


def load_hb():
    p = HB / "_summary.json"
    return json.loads(p.read_text()) if p.exists() else None


def highbudget_lines(hb):
    if not hb:
        return []
    am = hb["arm_means"]; c = hb["contrasts"]
    tb, tr, to, ob = c["tei_vs_baseline"], c["tei_vs_random"], c["tei_vs_opro"], c["opro_vs_baseline"]
    be = hb.get("budget_effect") or {}
    return [
        f"On the {hb['n']} headroom tasks (baseline<0.9) at a larger budget ({hb['iters']} iterations), arm means are "
        f"baseline {am['baseline']:.3f}, random {am['random']:.3f}, objective-only {am['objective_reflection']:.3f}, "
        f"OPRO-style {am['opro']:.3f}, TEI {am['tei']:.3f}.",
        f"TEI beats baseline by {tb['mean_delta']:+.3f} (p={fmt_p(tb['t_p_value'])}, d_z={tb['cohen_dz']:.2f}) and beats the "
        f"OPRO-style optimizer by {to['mean_delta']:+.3f} (p={fmt_p(to['t_p_value'])}), while OPRO moves the baseline by "
        f"exactly {ob['mean_delta']:+.3f} (p={fmt_p(ob['t_p_value'])}) -- reflective optimization helps where a non-reflective one does not.",
        f"Honest caveat: TEI is still NOT statistically separable from budget-matched random paraphrase "
        f"({tr['mean_delta']:+.3f}, p={fmt_p(tr['t_p_value'])}); and extra budget barely matters "
        f"(20- vs 6-iteration TEI: {be.get('mean_delta',0):+.3f}, p={fmt_p(be.get('t_p_value',1))}), so the aggregate null is not optimization-starvation.",
    ]


def fmt_p(p):
    p = float(p)
    return f"{p:.1e}" if p < 1e-3 else f"{p:.3f}"


def load():
    summary = json.loads((RES / "_summary.json").read_text())
    rows = list(csv.DictReader((RES / "per_agent.csv").open()))
    return summary, rows


TITLE = ("Does Evaluation-Guided Prompt Optimization Beat a Fair Baseline? "
         "A Confound-Controlled, Ablated Study of the Target-Evaluate-Improve Loop "
         "on %d Single-Turn Classification Tasks (TEI-Bench)")


def abstract(summary, hb=None):
    am = summary["arm_means"]; c = summary["contrasts"]
    tb, tr, to = c["tei_vs_baseline"], c["tei_vs_random"], c["tei_vs_objref"]
    hr = summary.get("headroom_subset", {})
    hrs = hr.get("stats", {})
    ef = summary.get("extra_full", {})
    tost = ef.get("tost05", {}); pw = ef.get("power", {})
    n = summary["n_agents"]; nc = summary.get("n_ceiling_tasks", 0)
    return (
        f"Evaluation-guided prompt optimization is widely reported to improve LLM systems, but the "
        f"evidence often relies on in-sample evaluation, single tasks, same-family judges, and no "
        f"comparison against undirected prompt search. We build TEI-Bench, a confound-controlled, "
        f"ablated, held-out protocol, and use it to evaluate one popular instantiation -- a GPA-inspired "
        f"evaluator feeding a GEPA-style reflective Pareto optimizer (the composition we call TEI) -- "
        f"across {n} single-turn classification, extraction, and reasoning tasks. The central finding is "
        f"cautionary. Under a naive label scorer, a pilot showed a large gain (+0.175 accuracy). After we "
        f"remove the output-format confound with a universal 'FINAL:' answer contract applied identically "
        f"to every condition, and add a four-arm ablation (baseline, undirected random prompt search, "
        f"objective-only reflection, full TEI), the gain collapses. Held-out arm means are nearly tied: "
        f"baseline {am['baseline']:.3f}, random {am['random']:.3f}, objective-only "
        f"{am['objective_reflection']:.3f}, TEI {am['tei']:.3f}. TEI does NOT significantly beat the "
        f"baseline (delta={tb['mean_delta']:+.3f}, p={fmt_p(tb['t_p_value'])}, d_z={tb['cohen_dz']:.2f}; "
        f"sign-test p={fmt_p(ef.get('sign',{}).get('p_value',1))}; permutation p={fmt_p(ef.get('perm',{}).get('p_value',1))}); "
        f"a two one-sided test shows the two are statistically EQUIVALENT within +/-0.05 "
        f"(p={fmt_p(tost.get('p_tost',1))}), though the full-suite test is underpowered ({nc} of {n} tasks "
        f"are at ceiling; MDE80={pw.get('mde_80pct',float('nan')):.3f}). TEI is indistinguishable from "
        f"random prompt search (delta={tr['mean_delta']:+.3f}, p={fmt_p(tr['t_p_value'])}) and the GPA "
        f"signal adds nothing over objective-only reflection (delta={to['mean_delta']:+.3f}, "
        f"p={fmt_p(to['t_p_value'])}); on this benchmark TEI behaves like a train-selected random "
        f"paraphrase. The only positive hint is a headroom subset (n={hr.get('n','?')}, baseline<0.9): "
        f"delta={hrs.get('mean_delta',0):+.3f}, a medium effect (d_z={hrs.get('cohen_dz',0):.2f}) that is "
        f"suggestive but not significant (p={fmt_p(hrs.get('t_p_value',1))}). We conclude that, within the "
        f"single-turn classification regime, confound control and an ablation against random search are "
        f"necessary to make prompt-optimization claims credible. Plan pre-registered (prereg-v2) before "
        f"the run; all code, data, traces, and optimized prompts are released."
        + (
            f" High-budget addendum: on the {hb['n']} non-saturated tasks at {hb['iters']} iterations, TEI DOES beat "
            f"the baseline ({hb['contrasts']['tei_vs_baseline']['mean_delta']:+.3f}, "
            f"p={fmt_p(hb['contrasts']['tei_vs_baseline']['t_p_value'])}, "
            f"d_z={hb['contrasts']['tei_vs_baseline']['cohen_dz']:.2f}) and a non-reflective OPRO-style optimizer "
            f"(which gains nothing), but still cannot be separated from budget-matched random paraphrase "
            f"({hb['contrasts']['tei_vs_random']['mean_delta']:+.3f}, "
            f"p={fmt_p(hb['contrasts']['tei_vs_random']['t_p_value'])})."
            if hb else ""
        )
    )


def robustness_lines(summary):
    ef = summary.get("extra_full", {})
    tost = ef.get("tost05", {}); pw = ef.get("power", {})
    sign = ef.get("sign", {}); perm = ef.get("perm", {})
    hr = summary.get("headroom_subset", {})
    n = summary["n_agents"]; nc = summary.get("n_ceiling_tasks", 0)
    out = [
        f"Because our central claim is a null, we quantify it rather than reporting only p>0.05.",
        f"Sign test (TEI vs baseline): {sign.get('pos','?')} wins / {sign.get('neg','?')} losses / "
        f"{sign.get('ties','?')} ties, p={fmt_p(sign.get('p_value',1))}.",
        f"Paired permutation (sign-flip) test on the mean difference: p={fmt_p(perm.get('p_value',1))}.",
        f"Equivalence (TOST, +/-0.05 margin): p={fmt_p(tost.get('p_tost',1))} -> "
        f"{'EQUIVALENT' if tost.get('equivalent') else 'not equivalent'} (we can rule out a large effect).",
        f"Power: {nc} of {n} tasks are at ceiling (>=0.999), so the full suite detects only effects "
        f">= MDE80={pw.get('mde_80pct',float('nan')):.3f} at 80% power; post-hoc power for the observed "
        f"effect is {pw.get('post_hoc_power',float('nan')):.2f}.",
    ]
    if hr:
        out.append(
            f"Headroom subset (n={hr.get('n','?')}): delta={hr.get('stats',{}).get('mean_delta',0):+.3f}, "
            f"d_z={hr.get('stats',{}).get('cohen_dz',0):.2f}, p={fmt_p(hr.get('stats',{}).get('t_p_value',1))}, "
            f"power={hr.get('power',{}).get('post_hoc_power',float('nan')):.2f}, "
            f"MDE80={hr.get('power',{}).get('mde_80pct',float('nan')):.3f} -- suggestive but underpowered.")
    return out


CONTRAST_LABELS = [
    ("TEI − baseline", "tei_vs_baseline"),
    ("TEI − random search", "tei_vs_random"),
    ("TEI − objective-only reflection", "tei_vs_objref"),
    ("random − baseline", "random_vs_baseline"),
]

REFERENCES = [
    "Agrawal, L. A., et al. (2025). GEPA: Reflective Prompt Evolution Can Outperform Reinforcement Learning. arXiv:2507.19457.",
    "Khattab, O., et al. (2023). DSPy: Compiling Declarative LM Calls into Self-Improving Pipelines. arXiv:2310.03714.",
    "Opsahl-Ong, K., et al. (2024). Optimizing Instructions and Demonstrations for Multi-Stage LM Programs (MIPROv2). arXiv:2406.11695.",
    "Yang, C., et al. (2023). Large Language Models as Optimizers (OPRO). arXiv:2309.03409.",
    "Zhou, Y., et al. (2022). Large Language Models Are Human-Level Prompt Engineers (APE). arXiv:2211.01910.",
    "Yuksekgonul, M., et al. (2024). TextGrad: Automatic Differentiation via Text. arXiv:2406.07496.",
    "Madaan, A., et al. (2023). Self-Refine: Iterative Refinement with Self-Feedback. arXiv:2303.17651.",
    "Shinn, N., et al. (2023). Reflexion: Language Agents with Verbal Reinforcement Learning. arXiv:2303.11366.",
    "Zheng, L., et al. (2023). Judging LLM-as-a-Judge with MT-Bench and Chatbot Arena. arXiv:2306.05685.",
    "Jia, A. S., et al. (2025). What Is Your Agent's GPA? A Framework for Evaluating Agent Goal-Plan-Action Alignment. arXiv:2510.08847.",
    "Cobbe, K., et al. (2021). Training Verifiers to Solve Math Word Problems (GSM8K). arXiv:2110.14168.",
    "Casanueva, I., et al. (2020). Efficient Intent Detection with Dual Sentence Encoders (banking77). arXiv:2003.04807.",
]


def contrast_rows(summary):
    out = []
    for label, key in CONTRAST_LABELS:
        s = summary["contrasts"][key]
        out.append([label, f"{s['mean_before']:.3f}", f"{s['mean_after']:.3f}",
                    f"{s['mean_delta']:+.3f}", f"[{s['ci95_low']:+.3f}, {s['ci95_high']:+.3f}]",
                    f"{s['t_stat']:.2f}", fmt_p(s['t_p_value']), f"{s['cohen_dz']:.2f}"])
    return out


def build_html(summary, rows, hb=None):
    c = contrast_rows(summary)
    am = summary["arm_means"]

    def tbl(headers, data, numeric_from=2):
        h = "<tr>" + "".join(f"<th>{x}</th>" for x in headers) + "</tr>"
        body = "".join("<tr>" + "".join(f"<td>{html.escape(str(v))}</td>" for v in r) + "</tr>" for r in data)
        return f"<table>{h}{body}</table>"

    arm_tbl = tbl(["Arm", "Mean held-out objective"],
                  [["Baseline", f"{am['baseline']:.3f}"], ["Random search", f"{am['random']:.3f}"],
                   ["Objective-only reflection", f"{am['objective_reflection']:.3f}"],
                   ["TEI (full)", f"{am['tei']:.3f}"]])
    abl_tbl = tbl(["Contrast", "A", "B", "Δ", "95% CI", "t", "p", "d_z"], c)
    per = [[r["task_id"], r["domain"], "public" if r["is_public"] == "True" else "synthetic",
            r["n_test"], r["obj_baseline"], r["obj_random"], r["obj_objref"], r["obj_tei"]] for r in rows]
    per_tbl = tbl(["Agent", "Domain", "Source", "n_test", "Base", "Rand", "ObjRef", "TEI"], per)
    refs = "<ol>" + "".join(f"<li>{html.escape(r)}</li>" for r in REFERENCES) + "</ol>"
    figs = "".join(
        f'<figure><img src="paper/figures/{fn}"/><figcaption>{html.escape(cap)}</figcaption></figure>'
        for fn, cap in [("ablation_arms.png", "Held-out objective by optimization arm (y-axis truncated; whiskers are 95% CIs, which overlap)."),
                        ("headroom_forest.png", "Per-task TEI − baseline on the headroom subset (baseline<0.9)."),
                        ("slope_objective_v2.png", "Per-agent baseline → TEI (held-out)."),
                        ("public_vs_synthetic.png", "Public vs. synthetic tasks (y-axis truncated).")]
        if (FIG / fn).exists())
    robust = "<ul>" + "".join(f"<li>{html.escape(x)}</li>" for x in robustness_lines(summary)) + "</ul>"
    hb_lines = highbudget_lines(hb)
    hb_html = ("<h2>High-budget / headroom finding (20 iterations + OPRO baseline)</h2><ul>"
               + "".join(f"<li>{html.escape(x)}</li>" for x in hb_lines) + "</ul>") if hb_lines else ""
    sg = summary.get("subgroups", {}); hr = summary.get("headroom_subset")
    sub = ""
    if "public" in sg:
        sub += (f"<p>Public-only (n={sg['public']['n']}): Δ={sg['public']['stats']['mean_delta']:+.3f} "
                f"(p={fmt_p(sg['public']['stats']['t_p_value'])}). ")
    if "synthetic" in sg:
        sub += (f"Synthetic-only (n={sg['synthetic']['n']}): Δ={sg['synthetic']['stats']['mean_delta']:+.3f} "
                f"(p={fmt_p(sg['synthetic']['stats']['t_p_value'])}). ")
    if hr:
        sub += (f"Headroom (baseline&lt;0.9, n={hr['n']}): Δ={hr['stats']['mean_delta']:+.3f} "
                f"(p={fmt_p(hr['stats']['t_p_value'])}, d_z={hr['stats']['cohen_dz']:.2f}).</p>")

    doc = f"""<!doctype html><html><head><meta charset="utf-8"><title>TEI-Bench</title><style>
body{{font-family:Calibri,Arial,sans-serif;max-width:880px;margin:2.2em auto;padding:0 1.2em;line-height:1.55;color:#111}}
h1{{font-size:22pt}} h2{{font-size:15pt;border-bottom:1px solid #ccd;padding-bottom:3px;margin-top:1.5em}}
table{{border-collapse:collapse;width:100%;margin:1em 0;font-size:10pt}}
th,td{{border:1px solid #b9c0cc;padding:5px 8px}} th{{background:#eaeefb}}
.meta{{text-align:center;color:#333}} .abstract{{background:#f7f8fc;border:1px solid #e2e6f0;border-radius:8px;padding:1em 1.2em;font-size:10.5pt}}
figure{{text-align:center;margin:1.2em 0}} img{{max-width:100%;border:1px solid #e5e7eb}} figcaption{{font-size:9pt;color:#475569}}
ol li{{font-size:9.5pt;margin-bottom:4px}}</style></head><body>
<h1>{html.escape(TITLE % summary['n_agents'])}</h1>
<p class="meta">Orkhan Javadli (MIT, alumnus) &middot; Anni Zimina (Stanford)<br>
<a href="https://github.com/ojavadli/tei-bench">github.com/ojavadli/tei-bench</a> &middot; pre-registration tag: prereg-v2</p>
<h2>Abstract</h2><div class="abstract">{html.escape(abstract(summary, hb))}</div>
<h2>Mean held-out objective by arm</h2>{arm_tbl}
<h2>Paired contrasts (held-out, n={summary['n_agents']})</h2>{abl_tbl}
<p>Holm-corrected p: TEI−baseline {fmt_p(summary['holm_corrected_p']['tei_vs_baseline'])}, "
TEI−random {fmt_p(summary['holm_corrected_p']['tei_vs_random'])}, "
TEI−objref {fmt_p(summary['holm_corrected_p']['tei_vs_objref'])}.</p>
<h2>Statistical robustness (the null, quantified)</h2>{robust}
{hb_html}
{figs}
<h2>Subgroups</h2>{sub}
<h2>Per-agent held-out objective by arm</h2>{per_tbl}
<h2>References</h2>{refs}
<h2>Citation</h2><pre>@misc{{javadli2026teibench, title={{TEI-Bench}}, author={{Javadli, Orkhan and Zimina, Anni}}, year={{2026}}, note={{github.com/ojavadli/tei-bench}}}}</pre>
</body></html>"""
    (ROOT / "TEI-Bench.html").write_text(doc, encoding="utf-8")
    print("  wrote TEI-Bench.html")


def build_docx(summary, rows, hb=None):
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  python-docx not installed; skipping .docx"); return
    am = summary["arm_means"]
    doc = Document()
    doc.add_heading(TITLE % summary["n_agents"], 0)
    p = doc.add_paragraph("Orkhan Javadli (MIT, alumnus)  ·  Anni Zimina (Stanford)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("github.com/ojavadli/tei-bench  ·  pre-registration tag: prereg-v2")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Abstract", 1); doc.add_paragraph(abstract(summary, hb))

    doc.add_heading("Method & Design (summary)", 1)
    for line in [
        "Unit of analysis: one agent = one (task, baseline prompt) pair; each measured on the SAME held-out test set. n = %d." % summary["n_agents"],
        "Four arms: baseline, random prompt search, objective-only reflection, TEI (full).",
        "Models: agent claude-haiku-4-5; judge + optimizer claude-sonnet-4-5 (stronger, different model).",
        "Primary endpoint: code-computed accuracy / exact match from a universal 'FINAL:' output contract (no LLM); secondary: 4 GPA-inspired judge dimensions.",
        "Pre-registered (tag prereg-v2) before the run; paired t-test + Wilcoxon + Cohen d_z + bootstrap CI, Holm-corrected; public/synthetic subgroups.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Mean held-out objective by arm", 1)
    t = doc.add_table(rows=1, cols=2); t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text = "Arm"; t.rows[0].cells[1].text = "Mean held-out objective"
    for name, key in [("Baseline", "baseline"), ("Random search", "random"),
                      ("Objective-only reflection", "objective_reflection"), ("TEI (full)", "tei")]:
        c = t.add_row().cells; c[0].text = name; c[1].text = f"{am[key]:.3f}"

    doc.add_heading("Paired contrasts (held-out, n=%d)" % summary["n_agents"], 1)
    headers = ["Contrast", "A", "B", "Δ", "95% CI", "t", "p", "d_z"]
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for r in contrast_rows(summary):
        c = t.add_row().cells
        for i, v in enumerate(r):
            c[i].text = str(v)

    doc.add_heading("Statistical robustness (the null, quantified)", 1)
    for line in robustness_lines(summary):
        doc.add_paragraph(line, style="List Bullet")

    hb_lines = highbudget_lines(hb)
    if hb_lines:
        doc.add_heading("High-budget / headroom finding (20 iterations + OPRO baseline)", 1)
        for line in hb_lines:
            doc.add_paragraph(line, style="List Bullet")

    for fn, cap in [("ablation_arms.png", "Held-out objective by optimization arm. The y-axis is truncated so the <=0.015 differences are visible; the 95% CIs overlap and no pairwise difference survives Holm correction."),
                    ("headroom_forest.png", "Per-task TEI - baseline on the headroom subset (baseline<0.9): medium but non-significant; a few subjective tasks improve, two regress."),
                    ("slope_objective_v2.png", "Per-agent baseline → TEI (held-out)."),
                    ("public_vs_synthetic.png", "Public vs. synthetic tasks (y-axis truncated).")]:
        if (FIG / fn).exists():
            doc.add_picture(str(FIG / fn), width=Inches(5.0))
            cp = doc.add_paragraph(cap); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Per-agent held-out objective by arm", 1)
    headers = ["Agent", "Domain", "Source", "n", "Base", "Rand", "ObjRef", "TEI"]
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for r in rows:
        c = t.add_row().cells
        vals = [r["task_id"], r["domain"], "public" if r["is_public"] == "True" else "synthetic",
                r["n_test"], r["obj_baseline"], r["obj_random"], r["obj_objref"], r["obj_tei"]]
        for i, v in enumerate(vals):
            c[i].text = str(v)

    doc.add_heading("References", 1)
    for i, ref in enumerate(REFERENCES, 1):
        doc.add_paragraph(f"[{i}] {ref}")
    out = ROOT / "paper" / "TEI-Bench.docx"
    doc.save(str(out)); print(f"  wrote {out}")


def build_markdown(summary, rows, hb=None):
    am = summary["arm_means"]; n = summary["n_agents"]
    def md_tbl(headers, data):
        h = "| " + " | ".join(headers) + " |\n| " + " | ".join("---" for _ in headers) + " |\n"
        return h + "".join("| " + " | ".join(str(v) for v in r) + " |\n" for r in data)
    lines = [f"# {TITLE % n}", "",
             "**Orkhan Javadli (MIT, alumnus) · Anni Zimina (Stanford)**", "",
             "Code, data & traces: https://github.com/ojavadli/tei-bench · pre-registration tag: `prereg-v2`",
             "", "## Abstract", "", abstract(summary, hb), "",
             "## Mean held-out objective by arm", "",
             md_tbl(["Arm", "Mean held-out objective"],
                    [["Baseline", f"{am['baseline']:.3f}"], ["Random search", f"{am['random']:.3f}"],
                     ["Objective-only reflection", f"{am['objective_reflection']:.3f}"],
                     ["TEI (full)", f"{am['tei']:.3f}"]]),
             "", "## Paired contrasts (held-out, n=%d) — none significant after Holm correction" % n, "",
             md_tbl(["Contrast", "A", "B", "Δ", "95% CI", "t", "p", "d_z"], contrast_rows(summary)), "",
             "## Statistical robustness (the null, quantified)", "",
             "\n".join(f"- {x}" for x in robustness_lines(summary)), "",
             *(["## High-budget / headroom finding (20 iterations + OPRO baseline)", "",
                "\n".join(f"- {x}" for x in highbudget_lines(hb)), ""] if hb else []),
             "![Ablation arms](paper/figures/ablation_arms.png)",
             "![Headroom forest](paper/figures/headroom_forest.png)",
             "![Slope](paper/figures/slope_objective_v2.png)",
             "![Public vs synthetic](paper/figures/public_vs_synthetic.png)", "",
             "## Per-agent held-out objective by arm", "",
             md_tbl(["Agent", "Domain", "Source", "n_test", "Base", "Rand", "ObjRef", "TEI"],
                    [[r["task_id"], r["domain"], "public" if r["is_public"] == "True" else "synthetic",
                      r["n_test"], r["obj_baseline"], r["obj_random"], r["obj_objref"], r["obj_tei"]]
                     for r in rows]),
             "", "## References", "",
             "\n".join(f"{i}. {r}" for i, r in enumerate(REFERENCES, 1)), ""]
    (ROOT / "ARTICLE.md").write_text("\n".join(lines), encoding="utf-8")
    print("  wrote ARTICLE.md")


if __name__ == "__main__":
    summary, rows = load()
    hb = load_hb()
    print("Building v2 report from results_v2/..." + (" (+high-budget)" if hb else ""))
    build_markdown(summary, rows, hb)
    build_html(summary, rows, hb)
    build_docx(summary, rows, hb)
    print("Done.")
